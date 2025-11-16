import os
import logging
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches  # Добавлен правильный импорт
from .file_utils import load_company_documents
from .database import load_chat, get_company_chats_dir  # Добавлен get_company_chats_dir

logger = logging.getLogger(__name__)

COMPANY_DOCS_DIR = "company_docs"

def _add_formatted_text(paragraph, text):
    """Вспомогательная функция для добавления форматированного текста"""
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if i % 2 == 1:  # Текст между ** должен быть жирным
                paragraph.add_run(part).bold = True
            else:
                paragraph.add_run(part)
    else:
        paragraph.add_run(text)

async def generate_ai_response_for_report(messages: list) -> str:
    """Вспомогательная функция для генерации отчета через AI"""
    import aiohttp
    import asyncio
    import json
    
    AI_PROXY_URL = "http://localhost:11434"
    
    ai_request = {
        "messages": messages,
        "model": "qwen2.5:0.5b-instruct",
        "temperature": 0.3,
        "max_tokens": 2000,
        "stream": False
    }
    
    try:
        async with aiohttp.ClientSession() as session:
            async with session.post(
                f"{AI_PROXY_URL}/chat",
                json=ai_request,
                timeout=5000
            ) as response:
                
                if response.status != 200:
                    error_text = await response.text()
                    logger.error(f"Ошибка ИИ прокси при генерации отчета: {error_text}")
                    raise Exception(f"Ошибка ИИ: {error_text}")
                
                data = await response.json()
                
                # ИЗВЛЕКАЕМ ТОЛЬКО ТЕКСТ ОТЧЕТА ИЗ ПОЛЯ 'message'
                if 'message' in data and 'content' in data['message']:
                    return data['message']['content']
                elif 'response' in data:
                    return data['response']
                else:
                    # Если структура ответа неизвестна, попробуем найти текст в других полях
                    logger.warning(f"Неожиданная структура ответа ИИ: {data}")
                    # Попробуем извлечь текст из любого строкового поля
                    for key, value in data.items():
                        if isinstance(value, str) and len(value) > 50:  # Берем достаточно длинные строки
                            return value
                    # Если ничего не нашли, вернем весь ответ как строку (на крайний случай)
                    return str(data)
    
    except asyncio.TimeoutError:
        logger.error("Таймаут при генерации отчета ИИ")
        raise Exception("Таймаут при генерации отчета")
    except Exception as e:
        logger.error(f"Ошибка при обращении к ИИ для отчета: {e}")
        raise Exception(f"Ошибка соединения с ИИ: {str(e)}")

async def generate_financial_report():
    """Генерация финансового отчета с помощью нейросети"""
    
    # Загружаем все документы компании
    company_documents = load_company_documents()
    
    if not company_documents:
        raise Exception("Нет документов компании для анализа")
    
    # УЛУЧШЕННЫЙ ПРОМТ для нейросети
    system_prompt = """
    Ты - профессиональный финансовый аналитик. На основе предоставленных документов компании сгенерируй подробный финансовый отчет.

    КРИТИЧЕСКИ ВАЖНЫЕ ТРЕБОВАНИЯ:
    1. АБСОЛЮТНАЯ ДОСТОВЕРНОСТЬ - используй ТОЛЬКО ТЕ ДАННЫЕ, которые есть в документах
    2. Не добавляй никакой информации, которой нет в предоставленных документах!
    3. Если каких-то данных нет в документах - не придумывай их!

    СТРУКТУРА ОТЧЕТА (соблюдай точно):
    # Финансовый отчет компании

    ## Основные показатели
    [здесь только фактические данные из документов]

    ## Доходы
    [конкретные цифры и источники из документов]

    ## Расходы
    [конкретные цифры и статьи расходов из документов]

    ## Финансовые итоги
    [расчеты на основе реальных данных]

    ## Анализ
    [только выводы на основе реальных данных]

    ФОРМАТИРОВАНИЕ ДЛЯ DOCX (используй точно):
    - Заголовки: # Заголовок 1, ## Заголовок 2
    - Списки: начинай строки с - 
    - Жирный текст: оборачивай в ** ** (например: **1 000 000 руб**)
    - Только фактические данные из документов!

    НЕ ДОБАВЛЯЙ никаких пояснений, инструкций, комментариев о форматировании или мета-информации!
    Генерируй ТОЛЬКО чистый отчет в указанной структуре.
    """

    # Формируем сообщения для нейросети
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Документы компании для анализа:\n\n{company_documents}\n\nСгенерируй финансовый отчет ТОЛЬКО на основе этих данных. Не добавляй ничего от себя!"}
    ]
    
    try:
        # Используем нейросеть для генерации отчета
        report_content = await generate_ai_response_for_report(messages)
        
        logger.info(f"Получен ответ от ИИ для отчета: {report_content[:200]}...")
        
        # Создаем DOCX документ из сгенерированного отчета
        doc = Document()
        
        # УЛУЧШЕННЫЙ ПАРСИНГ с обработкой форматирования
        lines = report_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    current_paragraph = None
                continue
                
            # Обрабатываем заголовки
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
                current_paragraph = None
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
                current_paragraph = None
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
                current_paragraph = None
            # Обрабатываем элементы списка
            elif line.startswith('- '):
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    # Завершаем предыдущий абзац и начинаем новый для списка
                    current_paragraph = doc.add_paragraph()
                
                list_text = line[2:]
                _add_formatted_text(current_paragraph, list_text)
                # Добавляем маркер списка
                current_paragraph.style = 'List Bullet'
                current_paragraph = None
            # Обрабатываем нумерованные списки
            elif re.match(r'^\d+\. ', line):
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    current_paragraph = doc.add_paragraph()
                
                list_text = re.sub(r'^\d+\. ', '', line)
                _add_formatted_text(current_paragraph, list_text)
                current_paragraph.style = 'List Number'
                current_paragraph = None
            # Обычный текст
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    # Если уже есть абзац, добавляем перенос строки
                    current_paragraph.add_run('\n')
                
                _add_formatted_text(current_paragraph, line)
        
        # Сохраняем отчет
        report_filename = f"financial_report_ai_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        report_path = os.path.join(COMPANY_DOCS_DIR, report_filename)
        doc.save(report_path)
        
        logger.info(f"Финансовый отчет успешно сохранен: {report_path}")
        
        return report_path, report_filename
        
    except Exception as e:
        logger.error(f"Ошибка генерации отчета нейросетью: {e}")
        raise Exception(f"Не удалось сгенерировать отчет: {str(e)}")
    
async def generate_chat_history_docx(company_id: str, chat_id: str):
    """Генерация DOCX файла с историей чата"""
    
    # Загружаем чат
    chat = load_chat(company_id, chat_id)
    if not chat:
        raise Exception("Чат не найден")
    
    # Создаем документ
    doc = Document()
    
    # Заголовок
    title = doc.add_heading(f'История чата: {chat.title}', 0)
    
    # Информация о чате
    doc.add_paragraph(f'Компания ID: {chat.company_id}')
    doc.add_paragraph(f'Тип чата: {chat.type}')
    doc.add_paragraph(f'Пользователь: {chat.user_id}')
    doc.add_paragraph(f'Дата экспорта: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph()
    
    # Добавляем историю сообщений
    doc.add_heading('История сообщений', level=1)
    
    for message in chat.messages:
        # Создаем параграф для сообщения
        message_para = doc.add_paragraph()
        
        # Добавляем отправителя и время
        sender_run = message_para.add_run(f"{'Вы' if message.role == 'user' else 'AI Ассистент'} ")
        sender_run.bold = True
        
        time_run = message_para.add_run(f"({datetime.fromisoformat(message.timestamp).strftime('%H:%M %d.%m.%Y')}): ")
        time_run.italic = True
        # Добавляем содержимое сообщения
        content_run = message_para.add_run(message.content)
        
        # Добавляем отступ между сообщениями
        doc.add_paragraph()
    
    # Сохраняем документ
    filename = f"chat_history_{chat_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    export_dir = os.path.join(get_company_chats_dir(id), "exports")
    os.makedirs(export_dir, exist_ok=True)
    file_path = os.path.join(export_dir, filename)
    
    doc.save(file_path)
    
    return file_path, filename