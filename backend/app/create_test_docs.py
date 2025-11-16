import os
import pandas as pd
from docx import Document
from pathlib import Path

def create_test_documents():
    """Создание тестовых документов компании"""
    
    # Создаем папку для документов
    docs_dir = Path("data/company_docs")
    docs_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. Создаем финансовый отчет в DOCX
    doc = Document()
    doc.add_heading('ФИНАНСОВЫЙ ОТЧЕТ КОМПАНИИ "ТЕХНОИННОВАЦИИ"', 0)
    doc.add_heading('за 2024 год', level=1)
    
    doc.add_heading('ДОХОДЫ:', level=2)
    doc.add_paragraph('• Продажа программного обеспечения: 15,000,000 руб.')
    doc.add_paragraph('• Техническая поддержка: 3,500,000 руб.')
    doc.add_paragraph('• Консалтинговые услуги: 2,200,000 руб.')
    doc.add_paragraph('• Обучение: 800,000 руб.')
    doc.add_paragraph('ИТОГО ДОХОДЫ: 21,500,000 руб.')
    
    doc.add_heading('РАСХОДЫ:', level=2)
    doc.add_paragraph('• Заработная плата: 8,000,000 руб.')
    doc.add_paragraph('• Аренда офиса: 1,200,000 руб.')
    doc.add_paragraph('• Маркетинг и реклама: 1,500,000 руб.')
    doc.add_paragraph('• Закупка оборудования: 2,000,000 руб.')
    doc.add_paragraph('• Налоги: 3,800,000 руб.')
    doc.add_paragraph('• Прочие расходы: 1,000,000 руб.')
    doc.add_paragraph('ИТОГО РАСХОДЫ: 17,500,000 руб.')
    
    doc.add_heading('ЧИСТАЯ ПРИБЫЛЬ: 4,000,000 руб.', level=2)
    doc.save(docs_dir / "financial_report_2024.docx")
    
    # 2. Создаем отчет по проектам в XLSX
    projects_data = {
        'Проект': ['Система AI Chat', 'Мобильное приложение', 'Веб-платформа', 'Облачная инфраструктура'],
        'Бюджет': [2500000, 1800000, 3200000, 4000000],
        'Факт': [2300000, 2100000, 3000000, 3800000],
        'Отклонение': [-200000, 300000, -200000, -200000]
    }
    df_projects = pd.DataFrame(projects_data)
    df_projects.to_excel(docs_dir / "projects_q1_2024.xlsx", index=False)
    
    # 3. Создаем бюджет на 2025 год в XLSX
    budget_data = {
        'Статья': ['Разработка', 'Маркетинг', 'Административные', 'Исследования'],
        'План 2025': [12000000, 3000000, 2500000, 4000000],
        'Факт 2024': [9500000, 2500000, 2200000, 3200000]
    }
    df_budget = pd.DataFrame(budget_data)
    df_budget.to_excel(docs_dir / "budget_2025.xlsx", index=False)
    
    # 4. Создаем политику компании в DOCX
    policy_doc = Document()
    policy_doc.add_heading('ПОЛИТИКА КОМПАНИИ "ТЕХНОИННОВАЦИИ"', 0)
    
    policy_doc.add_heading('1. МИССИЯ КОМПАНИИ', level=1)
    policy_doc.add_paragraph('Разработка инновационных ИИ-решений для бизнеса.')
    
    policy_doc.add_heading('2. ЦЕННОСТИ:', level=1)
    policy_doc.add_paragraph('• Качество превыше всего')
    policy_doc.add_paragraph('• Инновации в каждом проекте')
    policy_doc.add_paragraph('• Клиентоориентированность')
    policy_doc.add_paragraph('• Профессиональное развитие')
    
    policy_doc.add_heading('3. ФИНАНСОВАЯ ПОЛИТИКА:', level=1)
    policy_doc.add_paragraph('• Ежеквартальная отчетность')
    policy_doc.add_paragraph('• Бюджетное планирование на год')
    policy_doc.add_paragraph('• Контроль расходов в рамках бюджета')
    policy_doc.add_paragraph('• Инвестиции в R&D: 15% от прибыли')
    
    policy_doc.add_heading('4. ПЕРСОНАЛ:', level=1)
    policy_doc.add_paragraph('Штат компании: 45 сотрудников')
    policy_doc.add_paragraph('Средняя зарплата: 180,000 руб./мес.')
    policy_doc.save(docs_dir / "company_policy.docx")
    
    # 5. Создаем SVG с инфографикой
    svg_content = '''<svg width="400" height="300" xmlns="http://www.w3.org/2000/svg">
  <rect x="50" y="50" width="300" height="200" fill="#f0f8ff" stroke="#4a76a8" stroke-width="2"/>
  <text x="70" y="80" font-family="Arial" font-size="16" fill="#333">Статистика компании 2024</text>
  <text x="70" y="110" font-family="Arial" font-size="14" fill="#333">Доход: 21.5 млн руб.</text>
  <text x="70" y="130" font-family="Arial" font-size="14" fill="#333">Расход: 17.5 млн руб.</text>
  <text x="70" y="150" font-family="Arial" font-size="14" fill="#333">Прибыль: 4.0 млн руб.</text>
  <text x="70" y="170" font-family="Arial" font-size="14" fill="#333">Сотрудники: 45 чел.</text>
  <text x="70" y="190" font-family="Arial" font-size="14" fill="#333">Проекты: 4 активных</text>
  <text x="70" y="210" font-family="Arial" font-size="14" fill="#333">R&D: 15% от прибыли</text>
</svg>'''
    
    with open(docs_dir / "company_stats.svg", 'w', encoding='utf-8') as f:
        f.write(svg_content)
    
    # 6. Создаем текстовый файл с общей информацией
    company_info = '''КОМПАНИЯ: ООО "ТехноИнновации"
СФЕРА ДЕЯТЕЛЬНОСТИ: Разработка программного обеспечения, ИИ решения
ОСНОВАНА: 2020 год
ШТАТ: 45 сотрудников
ЛОКАЦИЯ: Москва, Россия

КЛЮЧЕВЫЕ ПОКАЗАТЕЛИ 2024:
- Выручка: 21,500,000 рублей
- Операционные расходы: 17,500,000 рублей
- Чистая прибыль: 4,000,000 рублей
- Рентабельность: 18.6%

ОСНОВНЫЕ КЛИЕНТЫ:
- Банковский сектор (40%)
- Ритейл (30%)
- Телекоммуникации (20%)
- Другие (10%)'''
    
    with open(docs_dir / "company_info.txt", 'w', encoding='utf-8') as f:
        f.write(company_info)
    
    print("✅ Тестовые документы созданы в папке data/company_docs/")
    print("📁 Список созданных файлов:")
    for file in docs_dir.iterdir():
        print(f"   - {file.name}")

if __name__ == "__main__":
    create_test_documents()